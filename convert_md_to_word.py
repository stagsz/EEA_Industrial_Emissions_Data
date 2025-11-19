from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Read the markdown file
with open('dump/Frankfurt_MHKW_Verification.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Split by lines
lines = content.split('\n')

i = 0
while i < len(lines):
    line = lines[i]

    # Skip empty lines
    if not line.strip():
        i += 1
        continue

    # H1 heading (# )
    if line.startswith('# ') and not line.startswith('## '):
        heading = line.replace('# ', '').strip()
        h = doc.add_heading(heading, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # H2 heading (## )
    elif line.startswith('## ') and not line.startswith('### '):
        heading = line.replace('## ', '').strip()
        h = doc.add_heading(heading, level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # H3 heading (### )
    elif line.startswith('### ') and not line.startswith('#### '):
        heading = line.replace('### ', '').strip()
        h = doc.add_heading(heading, level=3)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # H4 heading (#### )
    elif line.startswith('#### '):
        heading = line.replace('#### ', '').strip()
        h = doc.add_heading(heading, level=4)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Horizontal rule (---)
    elif line.strip() == '---':
        doc.add_paragraph('_' * 80)

    # Tables (start with |)
    elif line.strip().startswith('|'):
        # Collect all table lines
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        i -= 1  # Back up one since we'll increment at the end

        # Parse table
        if len(table_lines) >= 2:
            # First line is headers
            headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

            # Second line is separator (skip it)
            # Remaining lines are data
            data_rows = []
            for row_line in table_lines[2:]:
                cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                data_rows.append(cells)

            # Create table
            if data_rows:
                table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
                table.style = 'Light Grid Accent 1'

                # Add headers
                for col_idx, header in enumerate(headers):
                    cell = table.rows[0].cells[col_idx]
                    cell.text = header
                    # Bold header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Add data
                for row_idx, row_data in enumerate(data_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        table.rows[row_idx + 1].cells[col_idx].text = cell_text

    # Code blocks (```)
    elif line.strip().startswith('```'):
        code_lines = []
        i += 1  # Skip the opening ```
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1

        # Add code block
        code_para = doc.add_paragraph('\n'.join(code_lines))
        code_para.style = 'No Spacing'
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    # Bullet lists (- )
    elif line.strip().startswith('- '):
        text = line.strip()[2:]
        p = doc.add_paragraph(text, style='List Bullet')

    # Numbered lists (1. )
    elif re.match(r'^\d+\.\s', line.strip()):
        text = re.sub(r'^\d+\.\s', '', line.strip())
        p = doc.add_paragraph(text, style='List Number')

    # Bold emphasis (**text**)
    else:
        # Regular paragraph
        p = doc.add_paragraph()

        # Handle bold (**text**)
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                # Regular text
                p.add_run(part)

    i += 1

# Save the document
output_file = 'dump/Frankfurt_MHKW_Verification.docx'
doc.save(output_file)
print(f"Word document created: {output_file}")
